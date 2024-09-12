import streamlit as st
import pandas as pd
from datetime import datetime, timedelta
import matplotlib.pyplot as plt
import seaborn as sns
import os
from docx import Document
from docx.shared import Inches
from io import BytesIO

# Set the title of the web app
col1, col2 = st.columns([3, 1])

with col1:
    st.title("DRUG INVENTORY AND DRUG TRACKING SYSTEM")

with col2:
    st.image("logo.png", use_column_width=True)

# Display the current date
st.write(f"Date: {datetime.now().strftime('%Y-%m-%d')}")

# Define the file paths for storing data
order_file_path = "orders.csv"
hospital_file_path = "hospitals.csv"
data_folder = "data"
inventory_file = os.path.join(data_folder, "inventory_dataset.csv")
vendor_file = os.path.join(data_folder, "vendor_dataset.csv")
consumption_file = os.path.join(data_folder, "consumption_dataset.csv")

# Function to check if CSV files exist and create them if they don't
def check_and_create_csv(file_path, columns):
    if not os.path.exists(file_path):
        df = pd.DataFrame(columns=columns)
        df.to_csv(file_path, index=False)

# Check and create required CSV files
check_and_create_csv(order_file_path, ["Order ID", "Order From", "Order Received", "Drug Name", "Quantity", "Received", "Confirmed", "Packed", "Dispatched", "Delivered", "Final Status"])
check_and_create_csv(hospital_file_path, ["Hospital ID", "Hospital Name", "Place", "Address", "Phone", "Email"])

# Load existing hospital data from the CSV file
hospital_data = pd.read_csv(hospital_file_path)

# Load existing order data from the CSV file
order_data = pd.read_csv(order_file_path)

# Load inventory data
if os.path.exists(inventory_file):
    inventory_data = pd.read_csv(inventory_file)

# Function to save hospital data to the CSV file
def save_hospital_data(df):
    try:
        df.to_csv(hospital_file_path, index=False)
        st.success("Hospital data saved successfully!")
    except Exception as e:
        st.error(f"Error saving hospital data: {e}")

# Function to save order data to the CSV file
def save_order_data(df):
    try:
        df.to_csv(order_file_path, index=False)
        st.success("Order data saved successfully!")
        
    except Exception as e:
        st.error(f"Error saving order data: {e}")




#  File uploader for inventory data
# st.header("Upload Inventory Data")
# inventory_file = st.file_uploader("Upload Inventory CSV", type=["csv"])

if inventory_file:
    inventory_data = pd.read_csv(inventory_file)
    st.subheader("Inventory Data")
    st.dataframe(inventory_data)
    
    # Alert for medicines with low quantity or nearing expiry
    low_quantity = inventory_data[inventory_data["quantity"] < 50]
    nearing_expiry = inventory_data[pd.to_datetime(inventory_data["expiry_date"]) <= datetime.now() + timedelta(days=90)]

    # Adding custom CSS for styling
    st.markdown(
        """
        <style>
        .custom-alert {
            background-color: red;
            padding: 10px;
            margin-bottom: 20px;
            border-radius: 5px;
            color: black;
            font-weight: 100px;
        }
        </style>
        """,
        unsafe_allow_html=True
    )

    # Alert for medicines with low quantity
    if not low_quantity.empty:
        st.markdown('<div class="custom-alert">Some medicines have low stock levels!</div>', unsafe_allow_html=True)
        st.dataframe(low_quantity)

    # if not nearing_expiry.empty:
    #     st.warning("Some medicines are nearing expiry!")
    #     st.dataframe(nearing_expiry)

    # Adding custom CSS for styling
    st.markdown(
        """
        <style>
        .custom-alert {
            background-color: red;
            padding: 10px;
            margin-bottom: 20px;
            border-radius: 5px;
            color: black;
            font-weight: 100px;
        }
        </style>
        """,
        unsafe_allow_html=True
    )

    # Alert for medicines with low quantity
    if not nearing_expiry.empty:
        st.markdown('<div class="custom-alert">Some medicines are nearing expiry!</div>', unsafe_allow_html=True)
        st.dataframe(nearing_expiry)

    # Visualization for Inventory Data
    st.subheader("Inventory Visualization")
    fig, ax = plt.subplots()
    sns.scatterplot(x="name", y="quantity", data=inventory_data, ax=ax)
    plt.axhline(y=50, color='red', linestyle='--', label='Threshold')
    ax.set_title("Inventory Levels of Drugs")
    ax.set_ylabel("Quantity")
    ax.set_xlabel("Name")
    ax.tick_params(axis='x', rotation=90)

    if len(ax.get_legend_handles_labels()[0]) > 0:
        plt.legend()

    st.pyplot(fig)


if consumption_file:
    consumption_data = pd.read_csv(consumption_file)
    st.subheader("Consumption Data")
    st.dataframe(consumption_data)
    
    # Visualization for Consumption Data
    st.subheader("Consumption Visualization")
    fig, ax = plt.subplots()
    sns.barplot(x="name", y="usage", data=consumption_data, ax=ax)
    ax.set_title("Drug Consumption Levels")
    ax.set_ylabel("Usage")
    ax.set_xlabel("Name")
    ax.set_xticklabels(ax.get_xticklabels(), rotation=90)

  
    st.pyplot(fig)





# Order Tracking Section
st.header("Order Tracking")
# Function to handle order tracking and hospital details

def add_order_tracking():
    global order_data, hospital_data, inventory_data  # Access the global data
    
    st.subheader("Add or Update Order Tracking")

    selected_drugs = []
    drug_quantities = {}
    
    # Select or Add Order ID
    order_ids = order_data["Order ID"].unique().tolist()
    order_ids.append("Add New Order")
    selected_order_id = st.selectbox("Select Order ID", order_ids)

    hospital_selection = hospital_data["Hospital Name"].iloc[0] if not hospital_data.empty else None

    
    if selected_order_id == "Add New Order":
        order_id = st.text_input("Enter New Order ID")

        # Hospital selection
        hospitals = hospital_data["Hospital Name"].tolist()
        hospitals.append("Add New Hospital")
        hospital_selection = st.selectbox("Select Hospital", hospitals)

        if hospital_selection == "Add New Hospital":
            st.subheader("Add New Hospital Details")
            hospital_id = st.text_input("Hospital ID")
            new_hospital_name = st.text_input("Hospital Name")
            place = st.text_input("Place")
            address = st.text_area("Address")
            phone = st.text_input("Phone")
            email = st.text_input("Email")
            inventory_file = st.file_uploader("Upload Legal Document", type=["pdf"])
            if inventory_file:
                hosp_options = ["Pending", "Rejected", "Verified"]
                verification_status = st.selectbox("Verification", hosp_options)
                if verification_status == "Verified":
                    st.success("Hospital Verified")
                    if st.button("Add Hospital"):
                        if hospital_id and new_hospital_name:
                            if hospital_id not in hospital_data["Hospital ID"].values:
                                new_hospital = pd.DataFrame({
                                    "Hospital ID": [hospital_id],
                                    "Hospital Name": [new_hospital_name],
                                    "Place": [place],
                                    "Address": [address],
                                    "Phone": [phone],
                                    "Email": [email],
                                })
                                hospital_data = pd.concat([hospital_data, new_hospital], ignore_index=True)
                                save_hospital_data(hospital_data)
                            else:
                                st.warning("Hospital ID already exists!")
                        else:
                            st.warning("Please fill in both Hospital ID and Name!")
                elif verification_status == "Rejected":
                    st.error("Hospital Rejected")
            hospital_selection = new_hospital_name

        order_date = st.date_input("Order Received Date", datetime.now())

        # Select multiple drugs and quantities from inventory
        if not inventory_data.empty:
            selected_drugs = st.multiselect("Select Drug Names", inventory_data["name"].tolist())
            drug_quantities = {}
            for drug in selected_drugs:
                max_quantity = inventory_data.loc[inventory_data["name"] == drug, "quantity"].values[0]
                drug_quantities[drug] = st.number_input(f"Quantity for {drug}", min_value=1, max_value=int(max_quantity), value=1)
        else:
            st.error("Inventory data is empty. Please upload inventory data.")

    else:
        order_id = selected_order_id
        st.write(f"Order ID: {order_id}")

        order_row = order_data[order_data["Order ID"] == order_id]
        st.dataframe(order_row)
        display_order_status(order_id)

    status_options = ["Received", "Confirmed", "Packed", "Dispatched", "Delivered"]
    order_status = st.selectbox("Update Order Status", status_options)

    if st.button("Add/Update Order"):
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        
        # Update inventory quantities and alert for low stock or nearing expiry
        if selected_order_id == "Add New Order":
            for drug, quantity in drug_quantities.items():
                if order_id in order_data["Order ID"].values:
                    order_row = order_data[(order_data["Order ID"] == order_id) & (order_data["Drug Name"] == drug)]
                    if not order_row.empty:
                        order_data.loc[order_row.index, ["Order From", "Order Received", "Quantity"]] = [hospital_selection, order_date, quantity]
                        order_data.loc[order_row.index, order_status] = timestamp
                    else:
                        new_order = pd.DataFrame({
                            "Order ID": [order_id],
                            "Order From": [hospital_selection],
                            "Order Received": [order_date],
                            "Drug Name": [drug],
                            "Quantity": [quantity],
                            "Received": [timestamp if order_status == "Received" else ""],
                            "Confirmed": [timestamp if order_status == "Confirmed" else ""],
                            "Packed": [timestamp if order_status == "Packed" else ""],
                            "Dispatched": [timestamp if order_status == "Dispatched" else ""],
                            "Delivered": [timestamp if order_status == "Delivered" else ""],
                            "Final Status": [order_status]
                        })
                        order_data = pd.concat([order_data, new_order], ignore_index=True)
                else:
                    new_order = pd.DataFrame({
                        "Order ID": [order_id],
                        "Order From": [hospital_selection],
                        "Order Received": [order_date],
                        "Drug Name": [drug],
                        "Quantity": [quantity],
                        "Received": [timestamp if order_status == "Received" else ""],
                        "Confirmed": [timestamp if order_status == "Confirmed" else ""],
                        "Packed": [timestamp if order_status == "Packed" else ""],
                        "Dispatched": [timestamp if order_status == "Dispatched" else ""],
                        "Delivered": [timestamp if order_status == "Delivered" else ""],
                        "Final Status": [order_status]
                    })
                    order_data = pd.concat([order_data, new_order], ignore_index=True)

                # Reduce inventory quantity
                inventory_data.loc[inventory_data["name"] == drug, "quantity"] -= quantity
                generate_invoice_word(order_id, hospital_selection, selected_drugs, list(drug_quantities.values()))
        else:
            order_data.loc[order_data["Order ID"] == order_id, order_status] = timestamp
            order_data.loc[order_data["Order ID"] == order_id, "Final Status"] = order_status

        save_order_data(order_data)
        inventory_file = os.path.join(data_folder, "inventory_dataset.csv")
        inventory_data.to_csv(inventory_file, index=False)
        
        


def display_order_status(order_id):
            # Get the order details for the selected order ID
            order_row = order_data[order_data["Order ID"] == order_id].iloc[0]
            
            # Define the status steps
            status_steps = ["Received", "Confirmed", "Packed", "Dispatched", "Delivered"]
            
            # Display the order details
            st.write(f"Order ID: {order_id}")
            st.write(f"Order From: {order_row['Order From']}")
            st.write(f"Order Received Date: {order_row['Order Received']}")
            
            # Display the order status as checkpoints
            for status in status_steps:
                if pd.notnull(order_row[status]) and order_row[status]:
                    st.write(f"✅ {status} ({order_row[status]})")
                else:
                    st.write(f"🔲 {status}")


def generate_invoice_word(order_id, hospital_name, drugs, quantities):
    # Create a new Word Document
    doc = Document()
    
    # Add Title
    doc.add_heading('Invoice', 0)
    
    # Fetch hospital details
    hospital_info = hospital_data[hospital_data["Hospital Name"] == hospital_name].iloc[0]
    
    # Add Hospital Details
    doc.add_heading('Hospital Details', level=1)
    doc.add_paragraph(f"Hospital: {hospital_info['Hospital Name']}")
    doc.add_paragraph(f"Address: {hospital_info['Address']}")
    doc.add_paragraph(f"Phone: {hospital_info['Phone']}")
    doc.add_paragraph(f"Email: {hospital_info['Email']}")
    
    # Add Order Details
    doc.add_heading('Order Details', level=1)
    doc.add_paragraph(f"Order ID: {order_id}")
    doc.add_paragraph(f"Order Date: {datetime.now().strftime('%Y-%m-%d')}")
    
    # Add a table for drugs, quantities, and prices
    doc.add_heading('Drug Details', level=1)
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Drug Name'
    hdr_cells[1].text = 'Quantity'
    hdr_cells[2].text = 'Price/Unit'
    hdr_cells[3].text = 'Total Price'
    
    # Initialize total price
    total_price = 0
    
    # Pre-fetch price data to avoid repeated lookups
    prices = inventory_data.set_index("name")["price_per_unit"].to_dict()
    
    # Loop through the drugs to add data to the Word document
    for drug, quantity in zip(drugs, quantities):
        price = prices.get(drug, 0)  # Fetch price from dictionary
        total = quantity * price
        total_price += total
        
        row_cells = table.add_row().cells
        row_cells[0].text = drug
        row_cells[1].text = str(quantity)
        row_cells[2].text = f"${price:.2f}"
        row_cells[3].text = f"${total:.2f}"
    
    # Add Total Price
    doc.add_paragraph(f"Total Bill Value: ${total_price:.2f}", style='Heading2')
    
    # Convert the Word document to a byte stream
    doc_stream = BytesIO()
    doc.save(doc_stream)
    doc_stream.seek(0)
    
    # Display Word document download button
    st.success("Invoice generated successfully!")
    st.download_button(label="Download Invoice", data=doc_stream, file_name=f"invoice_{order_id}.docx", mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

# Call the function to add or update order tracking
add_order_tracking()
